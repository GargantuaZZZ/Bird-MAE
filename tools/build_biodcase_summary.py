from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "BioDCASE_2025_2026_动物声学比赛总结.docx"

NAVY = RGBColor(31, 78, 121)
TEAL = RGBColor(32, 115, 113)
INK = RGBColor(35, 39, 42)
MUTED = RGBColor(90, 98, 104)
LIGHT = "EAF1F5"
LIGHT2 = "F4F7F8"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color="B9C4CA", size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para(doc, text="", size=10.5, bold=False, color=INK, italic=False,
             align=None, before=0, after=6, line=1.2, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run_font(r, size=10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.5])
    set_table_borders(table, color="9BB8C4", size="5")
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT)
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_task(doc, title, purpose, data, metric, status, takeaways):
    add_heading(doc, title, 2)
    for label, value in (
        ("任务目标", purpose),
        ("数据与设置", data),
        ("核心指标", metric),
        ("结果/状态", status),
    ):
        p = add_para(doc, before=0, after=4, line=1.18)
        r = p.add_run(f"{label}：")
        set_run_font(r, size=10.3, bold=True, color=TEAL)
        r = p.add_run(value)
        set_run_font(r, size=10.3)
    for item in takeaways:
        add_bullet(doc, item)


def make_table(doc, headers, rows, widths, small=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        shade_cell(cell, "2E6673")
        set_cell_margins(cell, top=110, bottom=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=small, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        if row_idx % 2:
            for cell in cells:
                shade_cell(cell, LIGHT2)
        for idx, text in enumerate(values):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, size=small, bold=(idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 15, 7),
        ("Heading 2", 12.5, TEAL, 10, 5),
        ("Heading 3", 11.2, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(10.2)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("BioDCASE 2025–2026  |  动物声学比赛总结")
    set_run_font(r, size=8.3, color=MUTED)
    r = p.add_run("  |  ")
    set_run_font(r, size=8.3, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    add_footer(section)

    add_para(doc, "BIOACOUSTIC DATA CHALLENGES", size=10, bold=True, color=TEAL,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=12)
    add_para(doc, "BioDCASE 2025–2026", size=29, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "动物声学比赛总结与趋势分析", size=18, bold=False, color=INK,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_para(doc, "覆盖 2025 年已完成赛题与 2026 年进行中赛题",
             size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "资料状态：2026 年 6 月 15 日",
             size=10, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=38)
    add_callout(
        doc,
        "一句话结论",
        "BioDCASE 从 2025 年的“物种/事件识别基准”快速扩展到 2026 年的跨域泛化、零样本识别、主动学习、声景检索和声学—eDNA 跨模态预测，研究重心正从单一榜单精度转向可迁移、低标注成本和生态学可用性。",
    )
    add_para(doc, "本文依据 BioDCASE 官方 Challenge 页面及 2025 年结果页整理。2026 年最终结果尚未公布，文中仅报告官方基线与赛程，不将其表述为最终排名。",
             size=9.5, color=MUTED, italic=True, after=0)
    doc.add_page_break()

    add_heading(doc, "1. 执行摘要", 1)
    add_bullet(doc, "2025 年为首届 BioDCASE Challenge，共 3 个任务：蚊虫声音分类、少样本声事件检测、海洋声景事件检测；均已公布最终结果。")
    add_bullet(doc, "2026 年扩展为 6 个任务，覆盖野生动物跨域监测、鲸类强标注检测、鸟类零样本识别、主动学习、生态信息检索及声学—eDNA 跨模态预测。")
    add_bullet(doc, "2025 年结果显示：预训练音频基础模型、模型集成、多尺度特征和针对类不平衡/域偏移的训练策略是高分方案的共同要素。")
    add_bullet(doc, "2026 年的任务设计更贴近真实生态工作流：目标物种可能不在训练类表中；标注预算受限；训练和测试地点不同；最终输出要服务物种出现、群落组成或生态调查。")

    add_heading(doc, "两届总体对比", 2)
    make_table(
        doc,
        ["维度", "BioDCASE 2025", "BioDCASE 2026"],
        [
            ("任务数", "3", "6"),
            ("状态", "已结束，结果已发布", "进行中；提交截止 2026-06-29"),
            ("主要物种/场景", "蚊虫、跨物种少样本、海洋哺乳动物/鱼类声景", "鸟/蛙/蝙蝠、鲸、鸟类零样本、陆海主动学习、综合声景、eDNA"),
            ("核心问题", "分类、检测、少样本适应", "跨域泛化、开放词汇/零样本、标注效率、跨模态生态推断"),
            ("常见指标", "Macro-F1、事件 F1、PSDS", "mAP、IoU-F1、AULC、AUROC、Macro-F1/Jaccard"),
            ("方法趋势", "预训练模型 + 集成开始占优", "基础模型成为基础设施，进一步考察迁移与生态决策价值"),
        ],
        [1.05, 2.55, 2.90],
        small=8.7,
    )

    add_heading(doc, "2. BioDCASE 2025：首届比赛", 1)
    add_para(doc, "首届比赛集中检验三类基础能力：细粒度分类、极少标注条件下的时间定位、复杂长时声景中的多类事件检测。总体时间线为 2025 年 4 月开放，6 月底提交，7 月公布结果，10 月在悉尼举办 workshop。")

    add_task(
        doc,
        "Task 1  蚊虫分类（Mosquito Classification）",
        "从飞行声/嗡鸣声中识别蚊虫类别，并评估模型能否利用分类学层级信息。",
        "两个子任务：① 32 个蚊种的多类分类，并提供属/亚属/种的层级标签；② 蚊虫与非蚊虫二分类，采用弱标签。训练数据规模有限、类别不均衡，且录音设备和环境存在变化。",
        "子任务 1 使用 Macro-F1 与层级 F1；子任务 2 使用 F1，强调少数类和层级错误的影响。",
        "最终排名中，Task 1a 的 NTU-AI 方案以 77.9% 的物种级 Macro-F1 领先；Task 1b 的 NTU-AI 方案取得约 98.0% F1。高分系统普遍采用预训练音频模型、数据增强与模型集成。",
        [
            "分类学层级可以作为损失函数、标签平滑或层级解码的先验，而不只是赛后分析信息。",
            "设备差异与短录音条件使域鲁棒性和时频增强非常关键。",
            "对 Bird-MAE 类方法而言，可重点比较遮掩预训练表征与通用音频基础模型在小样本、层级分类上的差异。",
        ],
    )

    add_task(
        doc,
        "Task 2  少样本声事件检测（Few-shot Bioacoustic Event Detection）",
        "仅给出目标声型的极少正例，在长录音中定位同类事件，模拟生态学家快速检索新物种/新叫声的工作流。",
        "开发集由多个物种和不同录音条件组成。每段测试录音提供最初 5 个正例作为支持集，模型需要检测后续同类事件。比赛设置常规赛道与更强调跨类适应的赛道。",
        "事件级 F-measure，匹配阈值采用时间 IoU 0.3；重点同时约束漏检、误检和边界定位。",
        "最佳系统约为 59.1% F1，显著超过约 9.1% 的基线；领先方案使用预训练表征、原型/相似度匹配、伪标签或集成。结果也显示，不同物种与数据集之间的性能方差仍很大。",
        [
            "这是最接近“给几个样例就检索整片录音”的实用任务，适合研究提示式/原型式音频检测。",
            "检测性能不仅取决于表征，还强烈受阈值校准、后处理和事件合并规则影响。",
            "跨数据集方差说明统一模型仍需更好的域适应与不确定性估计。",
        ],
    )

    add_task(
        doc,
        "Task 3  海洋声景物种检测（Marine Soundscape Benchmark）",
        "在多地点、强噪声、长时海洋录音中检测鲸类等生物声事件，检验跨地点泛化。",
        "训练/验证数据来自多个海洋区域，包含不同水听器、背景噪声和目标声型；测试地点与训练条件存在明显分布偏移。",
        "主指标为 PSDS1，同时报告事件级 F1 等补充指标，以综合衡量不同阈值下的检测性能。",
        "2025 年结果中，基于预训练音频网络和集成的系统处于领先位置；官方结果页显示 BioLingual、EfficientAT 等表征/架构被高频采用。总体上，跨地点性能仍明显低于同域表现。",
        [
            "海洋任务的低采样率、超长时间结构与陆地鸟声并不相同，需要重新设计输入窗长和多尺度建模。",
            "PSDS 相比单阈值 F1 更能暴露校准和稳定性问题。",
            "预训练模型有效，但域外地点、设备噪声和稀有事件仍是主要瓶颈。",
        ],
    )

    add_heading(doc, "2025 年结果的共同规律", 2)
    make_table(
        doc,
        ["规律", "表现"],
        [
            ("基础模型迁移", "PANNs、BEATs、EfficientAT、BioLingual 等预训练表征成为高分系统常用起点。"),
            ("模型集成", "跨架构或跨折集成通常优于单模型，尤其在类别不均衡和域偏移条件下。"),
            ("数据增强", "Mixup、SpecAugment、时移、频率遮掩和噪声增强被用于缓解小数据问题。"),
            ("后处理重要", "事件阈值、平滑、合并和持续时间约束对检测榜单影响显著。"),
            ("泛化仍未解决", "地点、设备、物种和声景变化导致各子数据集性能差异很大。"),
        ],
        [1.35, 5.15],
        small=9.1,
    )

    doc.add_page_break()
    add_heading(doc, "3. BioDCASE 2026：任务扩展", 1)
    add_callout(
        doc,
        "当前状态",
        "截至 2026 年 6 月 15 日，比赛仍在进行。官方日程为：2026 年 3 月 30 日开放，6 月 29 日提交截止，7 月 6 日公布结果，9 月 22 日在都灵举行 workshop。因此下文的数值是任务设置或官方基线，不是最终获奖成绩。",
    )

    add_task(
        doc,
        "Task 1  鲁棒野生动物声学监测",
        "构建可跨物种群、地点和录音条件迁移的检测/分类系统。",
        "包含青蛙、鸟类和蝙蝠三个子任务。允许在官方指定的大规模数据上预训练，再在目标数据上微调或直接迁移；评价集强调地点与声学条件变化。",
        "以各子任务的 Macro-F1/mAP 等指标衡量，并关注跨域综合表现。",
        "进行中。该任务把“一个模型覆盖多类群”的基础模型能力放到真实监测数据上检验。",
        [
            "与 Bird-MAE 最直接相关：可比较专用鸟声预训练、通用音频预训练与多类群联合预训练。",
            "频率范围跨越蛙、鸟和超声蝙蝠，统一前端可能不是最优，多分辨率或分支式编码值得研究。",
        ],
    )

    add_task(
        doc,
        "Task 2  强标注鲸叫监督检测",
        "检测南极蓝鲸和长须鲸的 7 类叫声，并准确定位开始/结束时间。",
        "开发集包含 11 个站点—年份数据集、6591 个文件、约 1880 小时录音，采样率 250 Hz；训练与验证按地点/年份拆分。目标事件仅约占时长的 6%，评价集来自新地点/时期。",
        "时间轴 1D IoU 匹配后的分类 F1；重复检测同一个真值会被计为额外假阳性。",
        "这是 2025 同类任务的第二版，2026 修复了数据不一致、补充元数据，并重新公平计算 YOLO 基线。最终排名待公布。",
        [
            "极低采样率和长时上下文使图像式检测器与序列检测器都具有竞争空间。",
            "评估明确惩罚重复框，后处理中的去重和事件合并不可忽略。",
        ],
    )

    add_task(
        doc,
        "Task 3  鸟类零样本识别",
        "识别训练标签之外的鸟种，即目标类别在训练时不可见。",
        "结合鸟声录音与物种文本/分类学信息，测试模型能否把音频表征对齐到开放类别描述；核心难点是新物种、不同地理区域和长尾类别。",
        "以类别级识别指标（如 Macro-F1/mAP）评价零样本泛化。",
        "进行中。这一任务标志着比赛从封闭类表分类转向开放词汇生物声学。",
        [
            "音频—文本对齐、分类学层级和物种属性描述将成为关键监督信号。",
            "只优化已见类别准确率可能损伤未见类，需关注广义零样本校准。",
        ],
    )

    add_task(
        doc,
        "Task 4  动物声学主动学习",
        "在固定人工标注预算下，选择最值得标注的样本，使模型性能提升最快。",
        "使用预生成 Perch v2 的 5 秒嵌入，覆盖陆地 BirdSet 与海洋 ATBFL。每个子集最多选择 500 个样本，约为样本池的 5%；分类头和主动学习主循环固定，参赛者只设计采样函数。",
        "主指标为 mAP（macro）学习曲线下面积 AULC；补充报告采样计算成本、训练成本与实际标注成本。",
        "官方开发基线中 CoreSet 的聚合 AULC 约 0.422，高于随机采样约 0.401；最终结果待公布。",
        [
            "评价不再只看最终精度，而看每增加一批标注带来的收益，直接对应生态项目的人力预算。",
            "跨 BirdSet 与鲸类数据同时评估，能揭示采样策略是否真正跨域。",
        ],
    )

    add_task(
        doc,
        "Task 5  声景生态信息检索",
        "从声景中预测具有生态意义的多标签信息，如类群、栖息地或场景属性，而非只识别单一物种事件。",
        "数据来自自然声景，目标标签具有层级、多标签和不均衡特征；鼓励使用音频基础模型和元数据。",
        "采用宏平均 AUROC 等多标签指标，强调跨类别均衡与检索能力。",
        "进行中。该任务把音频表征从“检测器输入”提升为生态信息检索接口。",
        [
            "适合评估 Bird-MAE 表征是否保留了物种之外的环境、群落和声景信息。",
            "需要防止模型依赖地点或设备捷径，建议做地点隔离与元数据消融。",
        ],
    )

    add_task(
        doc,
        "Task 6  声学—eDNA 跨模态预测",
        "利用被动声学监测数据预测 eDNA 检出的物种组成，连接两种互补的生物多样性调查手段。",
        "任务包含三个层次：音频分类、按样点聚合的出现概率预测，以及完整音景到 eDNA 物种集合的端到端预测。数据涉及鸟类、蝙蝠、蛙类、昆虫和非发声类群，并提供地点/时间元数据。",
        "主要采用 Macro-F1 与 Jaccard 相似度，分别衡量类别均衡性能和预测物种集合重合度。",
        "进行中。官方给出了基于 Perch 等嵌入和环境/时空信息的基线，最终排名待公布。",
        [
            "这是最具生态学创新性的任务：目标不局限于“声音是谁发的”，而是推断当地群落。",
            "非发声物种只能通过共现、环境和时空关系间接预测，必须严防空间泄漏并解释模型依据。",
        ],
    )

    add_heading(doc, "4. 2025 → 2026 的研究趋势", 1)
    trends = [
        ("从封闭集到开放集", "2025 多数任务的目标类表固定；2026 加入鸟类零样本与跨模态群落预测，要求模型处理未知类别和间接生态信号。"),
        ("从模型精度到数据效率", "主动学习用 AULC 与标注成本衡量“少标多少数据还能学好”，更贴近项目预算。"),
        ("从单域到跨域", "同一任务跨鸟类、鲸类、蛙类、蝙蝠和昆虫，基础模型必须适应频段、时长和设备差异。"),
        ("从事件识别到生态推断", "声景检索与 eDNA 任务把输出从声事件拓展到栖息地、群落和生物多样性指标。"),
        ("从离线模型到工作流", "采样策略、阈值校准、后处理、计算成本和可复现提交成为评价的一部分。"),
    ]
    make_table(doc, ["趋势", "含义"], trends, [1.55, 4.95], small=9.1)

    add_heading(doc, "5. 对 Bird-MAE 研究的启示", 1)
    add_bullet(doc, "优先切入 2026 Task 1、Task 3 和 Task 5：它们分别检验跨域迁移、零样本语义对齐与声景级生态信息保留，最能体现自监督预训练价值。")
    add_bullet(doc, "不要只做线性探测准确率。建议同时报告少样本曲线、地点外推、设备外推、类别长尾、阈值稳定性和校准误差。")
    add_bullet(doc, "设计多时间尺度评估：短窗适合叫声片段，长窗适合声景与群落；可比较平均池化、注意力池化和层级时序建模。")
    add_bullet(doc, "加入分类学或文本信息进行音频—语义对齐，可自然延伸到 2026 鸟类零样本任务。")
    add_bullet(doc, "主动学习可直接使用 MAE 表征做多样性采样、原型覆盖或不确定性—多样性联合选择，验证表征对标注效率的贡献。")
    add_bullet(doc, "所有跨地点实验应按站点/年份拆分，避免相邻片段或同一录音泄漏到训练和测试。")

    add_heading(doc, "建议的实验矩阵", 2)
    make_table(
        doc,
        ["研究问题", "建议对照", "关键指标"],
        [
            ("预训练是否提升小样本？", "随机初始化 / ImageNet / 通用音频模型 / Bird-MAE", "1/5/10-shot F1、置信区间"),
            ("是否跨地点泛化？", "同地点划分 vs 留一地点测试", "Macro-F1、mAP、性能下降幅度"),
            ("是否支持未知物种？", "线性分类 vs 音频—文本对齐", "Seen/Unseen/Harmonic mean"),
            ("是否降低标注成本？", "Random / Uncertainty / CoreSet / MAE-based", "AULC、采样耗时、标注成本"),
            ("是否保留声景信息？", "片段池化 / 多尺度池化 / 层级时序模型", "Macro-AUROC、Jaccard"),
        ],
        [2.15, 2.80, 1.55],
        small=8.6,
    )

    add_heading(doc, "6. 结论", 1)
    add_para(doc, "BioDCASE 2025 建立了动物声学中分类、少样本检测和复杂声景检测的统一竞赛基准；2026 则迅速把问题推进到跨域基础模型、开放类别、主动标注和跨模态生态调查。对于研究团队，最值得关注的不是某个单一榜单架构，而是表征能否在新地点、新物种、低标注预算和生态学下游任务中稳定复用。")
    add_para(doc, "如果以 Bird-MAE 为主线，建议将“自监督表征 + 多尺度时序 + 分类学/文本对齐 + 标注效率”作为统一研究框架，并用 2025 已完成任务做可比基准、用 2026 任务验证前沿泛化能力。")

    add_heading(doc, "7. 主要资料来源", 1)
    sources = [
        ("BioDCASE Challenge 2025 总览", "https://biodcase.github.io/challenge2025/"),
        ("2025 Task 1 及结果", "https://biodcase.github.io/challenge2025/task1"),
        ("2025 Task 2 及结果", "https://biodcase.github.io/challenge2025/task2"),
        ("2025 Task 3 及结果", "https://biodcase.github.io/challenge2025/task3"),
        ("BioDCASE Challenge 2026 总览", "https://biodcase.github.io/challenge2026/"),
        ("2026 Task 1", "https://biodcase.github.io/challenge2026/task1"),
        ("2026 Task 2", "https://biodcase.github.io/challenge2026/task2"),
        ("2026 Task 3", "https://biodcase.github.io/challenge2026/task3"),
        ("2026 Task 4", "https://biodcase.github.io/challenge2026/task4"),
        ("2026 Task 5", "https://biodcase.github.io/challenge2026/task5"),
        ("2026 Task 6", "https://biodcase.github.io/challenge2026/task6"),
        ("Workshop 2025", "https://biodcase.github.io/workshop2025/"),
        ("Workshop 2026", "https://biodcase.github.io/workshop2026/"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(3)
        add_hyperlink(p, label, url)

    add_para(doc, "注：结果数字按官方页面展示进行概括；不同赛道可能同时报告主指标、次指标和多个提交版本，本文侧重研究结论而非逐项复刻完整排行榜。",
             size=9.0, italic=True, color=MUTED, before=8, after=0)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    doc.core_properties.title = "BioDCASE 2025–2026 动物声学比赛总结与趋势分析"
    doc.core_properties.subject = "动物声学竞赛、任务总结、研究趋势"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
