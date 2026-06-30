from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
IMG = ROOT / "assets" / "biodcase_images"
OUT = ROOT / "BioDCASE_2025_2026_动物声学比赛总结_更新版.docx"

BLUE = RGBColor(30, 75, 116)
TEAL = RGBColor(25, 115, 110)
INK = RGBColor(35, 39, 42)
MUTED = RGBColor(86, 96, 104)
WHITE = RGBColor(255, 255, 255)
FILL = "EAF2F4"
FILL2 = "F5F7F8"


def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    cant.set(qn("w:val"), "true")
    tr_pr.append(cant)


def borders(table, color="B8C6CD"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "5")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        tbl_borders.append(e)


def table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW")) or OxmlElement("w:tcW")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=10.5, bold=False, color=INK, italic=False,
             before=0, after=6, align=None, line=1.18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.36)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), size=10.1)


def add_link(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def make_table(doc, headers, rows, widths, small=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_widths(table, widths)
    borders(table)
    no_split(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "2E6673")
        cell_margin(cell, top=110, bottom=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(h), size=small, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        tr = table.add_row()
        no_split(tr)
        for i, value in enumerate(row):
            cell = tr.cells[i]
            if ridx % 2:
                shade(cell, FILL2)
            cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            font(p.add_run(str(value)), size=small, bold=(i == 0))
    add_para(doc, "", after=2)
    return table


def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    table_widths(t, [6.5])
    borders(t, color="8AA9B4")
    c = t.cell(0, 0)
    shade(c, FILL)
    cell_margin(c, top=130, bottom=130, start=170, end=170)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), size=10.5, bold=True, color=BLUE)
    font(p.add_run(text), size=10.5, color=INK)
    add_para(doc, "", after=2)


def add_image(doc, filename, caption, width=5.8):
    path = IMG / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = add_para(doc, caption, size=8.6, color=MUTED, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    return cap


def task_block(doc, year, task, title, focus, data, metric, status, image=None, caption=None, image_width=5.7):
    add_heading(doc, f"{year} {task}  {title}", 2)
    for label, val in [
        ("任务目标", focus),
        ("数据/设置", data),
        ("评价指标", metric),
        ("结果/状态", status),
    ]:
        p = add_para(doc, after=3, line=1.15)
        font(p.add_run(f"{label}："), size=10.2, bold=True, color=TEAL)
        font(p.add_run(val), size=10.2)
    if image:
        add_image(doc, image, caption, width=image_width)


def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 15, 7),
        ("Heading 2", 12.3, TEAL, 10, 4),
        ("Heading 3", 11.2, BLUE, 8, 3),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("BioDCASE 2025-2026 | 更新版 | "), size=8.3, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def build():
    doc = Document()
    configure(doc)

    add_para(doc, "BIOACOUSTIC DATA CHALLENGES", size=10, bold=True, color=TEAL,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=12)
    add_para(doc, "BioDCASE 2025-2026", size=28, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "动物声学比赛任务清单、图示与虎豹方向补充", size=17,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_para(doc, "资料状态：2026 年 6 月 22 日；2026 Challenge 提交截止日为 2026 年 6 月 22 日，结果计划 2026 年 7 月 6 日公布。",
             size=9.6, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    callout(
        doc,
        "这版主要修正",
        "上一版把部分赛题概括成了更泛化的动物声学任务，甚至误把 2026 Task 6 写成 eDNA。新版按官方 challenge 页面逐项列出：2025 年 3 个任务，2026 年 6 个任务，并补入官方图示和“虎/豹个体识别”方向。",
    )
    doc.add_page_break()

    add_heading(doc, "1. 每年有哪些比赛", 1)
    make_table(
        doc,
        ["年份", "Task", "官方题目", "一句话说明"],
        [
            ("2025", "Task 1", "Multi-Channel Alignment", "校正多录音设备之间因时钟漂移造成的时间不同步。"),
            ("2025", "Task 2", "Supervised Detection of Strongly-Labelled Whale Calls", "检测南极蓝鲸/长须鲸多类叫声，并输出强标注时间边界。"),
            ("2025", "Task 3", "Bioacoustics for Tiny Hardware", "在 ESP32-S3-Korvo-2 微控制器上做鸟声检测/识别，兼顾精度与资源。"),
            ("2026", "Task 1", "Multichannel Alignment", "延续并扩展多通道对齐，包含 ARU 与 zebra finch 数据。"),
            ("2026", "Task 2", "Supervised Detection of Strongly-Labelled Whale Calls", "延续鲸声强标注检测，继续关注稀疏事件和跨声景泛化。"),
            ("2026", "Task 3", "Bioacoustics for Tiny Hardware", "继续考察低功耗边缘设备上的鸟声模型。"),
            ("2026", "Task 4", "Active Learning for Bioacoustics", "在固定标注预算下选择最值得标注的样本，覆盖陆地鸟声与海洋鲸声。"),
            ("2026", "Task 5", "Cross-Domain Mosquito Species Classification", "在不同地点、设备和声环境之间泛化的蚊种分类。"),
            ("2026", "Task 6", "Bird Counting", "从动物园鸟舍声景中估计目标鸟种个体数量。"),
        ],
        [0.65, 0.65, 2.25, 2.95],
        small=8.4,
    )

    add_heading(doc, "2. 2025 赛题详情", 1)
    task_block(
        doc, "2025", "Task 1", "Multi-Channel Alignment",
        "给定成对的不同步录音，估计两路音频之间的时间映射，恢复同步。",
        "每个 stereo audio file 的两路通道因非线性 clock drift 而不同步；开发阶段提供少量人工 keypoints，评测阶段要求同步未见音频对。",
        "以同步误差为核心，官方页面报告 MAE/MSE 等误差；误差越小越好。",
        "2025 已结束并发布结果。该任务更偏“声学基础处理/阵列或多设备预处理”，不是物种分类。",
        "task1_keypoints.jpg",
        "官方图示：多通道对齐中的 keypoints，用来描述两路录音的时间对应关系。来源：BioDCASE 2025 Task 1。",
    )
    task_block(
        doc, "2025", "Task 2", "Supervised Detection of Strongly-Labelled Whale Calls",
        "在长时海洋被动声学监测数据中检测 7 类南极蓝鲸和长须鲸叫声，输出类别与起止时间。",
        "基于 AcousticTrends Blue/Fin Library；事件稀疏、背景声景变化大，且一个音频片段中可有多个事件。",
        "强标注 sound event detection 指标，关注事件类别和时间定位；叫声类别按评估组聚合。",
        "2025 已结束并发布结果。任务价值在于稀有事件检测、跨地点/跨年份声景泛化和长时海洋 PAM 自动化。",
        "task2_whale.png",
        "官方图示：蓝鲸/长须鲸叫声的强标注检测示例。来源：BioDCASE 2025 Task 2。",
    )
    task_block(
        doc, "2025", "Task 3", "Bioacoustics for Tiny Hardware",
        "在资源受限微控制器上部署鸟声模型，以低成本、低功耗设备完成边缘端生物声学分析。",
        "目标场景是 Yellowhammer 鸟声检测/识别；硬件平台为 ESP32-S3-Korvo-2，强调模型大小、推理时延和峰值内存。",
        "同时看分类性能和资源效率；需要报告模型大小、推理时间、峰值内存等。",
        "2025 已结束并发布结果。该任务和 Bird-MAE 的关系在于：预训练表征若要落地，需要蒸馏、剪枝或 tiny model 适配。",
        "task3_dev_board.jpg",
        "官方图示：ESP32-S3-Korvo-2，BioDCASE Tiny Hardware 任务使用的开发板。来源：BioDCASE 2025 Task 3。",
    )

    add_heading(doc, "3. 2026 赛题详情", 1)
    callout(doc, "赛程状态", "截至 2026 年 6 月 22 日，2026 Challenge 正处于提交截止日；官方结果计划 2026 年 7 月 6 日公布，BioDCASE Workshop 日期为 2026 年 10 月 23 日。")
    task_block(
        doc, "2026", "Task 1", "Multichannel Alignment",
        "继续解决多设备录音不同步问题，但加入更明确的域偏移测试。",
        "包含 aru 与 zebra_finch 两个数据集；train/val 有音频和 keypoints，test 只有音频。aru 为细粒度漂移，总不同步不超过约 +/-0.5 秒；zebra_finch 为粗粒度漂移，总不同步不超过约 +/-5 秒。",
        "MAE(ms) 与 MSE(sec^2) 等同步误差；越小越好。",
        "2026 进行中。官网基线包括 nosync、GCC-PHAT 和 deep learning baseline；deep learning 在 zebra_finch 粗漂移场景下更有优势。",
        "task1_keypoints.jpg",
        "官方图示沿用 keypoints 对齐概念。来源：BioDCASE 2026 Task 1 / 2025 Task 1。",
    )
    task_block(
        doc, "2026", "Task 2", "Supervised Detection of Strongly-Labelled Whale Calls",
        "延续鲸声强标注检测，目标仍是南极蓝鲸和长须鲸叫声的类别与时间边界。",
        "面向 marine PAM；核心困难包括目标叫声稀少、声景变化和跨地点/仪器泛化。",
        "强标注检测指标，关注时间边界和类别；需要避免重复检测同一事件。",
        "2026 进行中。该任务可作为长时低频动物声检测的代表性 benchmark。",
        "task2_whale.png",
        "官方图示：鲸声事件检测示意。来源：BioDCASE Task 2。",
    )
    task_block(
        doc, "2026", "Task 3", "Bioacoustics for Tiny Hardware",
        "继续推动鸟声模型在低成本边缘设备上运行。",
        "使用微控制器约束下的鸟声数据与 baseline framework，目标是在可部署硬件上达到可接受性能。",
        "分类性能与资源效率共同决定系统价值。",
        "2026 进行中。对 Bird-MAE 来说，重点不是直接上大模型，而是用预训练教师模型产生小模型、轻量特征或蒸馏方案。",
        "task3_dev_board.jpg",
        "官方图示：Tiny Hardware 开发板。来源：BioDCASE Task 3。",
    )
    task_block(
        doc, "2026", "Task 4", "Active Learning for Bioacoustics",
        "在固定专家标注预算下，设计 acquisition function，挑选最值得标注的音频片段。",
        "输入是预生成的 5 秒 Perch v2 embeddings；覆盖 BirdSet 与 ATBFL 两个域。每个子集最多选择 500 个样本。",
        "主指标为 mAP(macro) 学习曲线下面积 AULC；还报告训练成本、采样计算成本和标注成本。",
        "2026 进行中。官网基线显示 CoreSet 是强 baseline；任务强调方法必须跨陆地/海洋数据集泛化。",
        "task4_active_learning_loop.png",
        "官方图示：主动学习循环。来源：BioDCASE 2026 Task 4。",
    )
    task_block(
        doc, "2026", "Task 5", "Cross-Domain Mosquito Species Classification",
        "在录音条件变化时识别蚊虫物种，主目标是跨域泛化。",
        "开发集包含 5 个 domain、9 个 species、271,380 个 clips，总时长约 60.66 小时；domain 与 species 分布高度不均衡。",
        "BA_unseen 是主排名指标；DSG=|BA_unseen-BA_seen| 是次指标，越小代表 seen/unseen domain 差距越小。",
        "2026 进行中。官方 baseline 使用 8 kHz 音频、64-bin log-mel 和轻量 MTRCNN，并含 species head 与 domain head。",
        "task5_mosquito_icon.png",
        "官方图标：Cross-Domain Mosquito Species Classification。来源：BioDCASE 2026 Task 5。",
        1.35,
    )
    task_block(
        doc, "2026", "Task 6", "Bird Counting",
        "从动物园鸟舍声景中估计目标鸟种个体数量，而不是只做物种是否出现。",
        "开发集来自 6 个 aviaries，目标物种包括 Greater flamingo、Hadada ibis、Red-billed quelea；评测集为 held-out aviaries，并可选扩展到 Pied avocet 与完整 67 物种 inventory。",
        "主指标为 MAE，越低越好；补充报告 RMSE、R2、MAPE。baseline 是 species detection 后接 detection-count regression。",
        "2026 进行中。官方 baseline 使用 ARIA/BirdNET 检测、80+ detection/acoustic features，再做数量估计。",
        "task6_bird_icon.png",
        "官方图标：Bird Counting。来源：BioDCASE 2026 Task 6。",
        1.8,
    )

    add_heading(doc, "4. 和虎/豹声音方向的连接", 1)
    add_image(doc, "amur_tiger_eye.jpg", "照片候选：东北虎/西伯利亚虎。来源：Wikimedia Commons，Amur Tiger Panthera tigris altaica Eye。", width=3.1)
    add_image(doc, "amur_leopard_camera_trap.jpg", "照片候选：东北豹/远东豹相机陷阱画面。来源：Wikimedia Commons，Amur leopard camera trap frame。", width=4.7)
    doc.add_page_break()
    callout(
        doc,
        "方向定义",
        "虎豹方向不应写成“识别是哪一种虎/豹”。在保护监测里更自然的任务是同一物种内的个体识别：给定 roar、call、chuff 或其他可用声型，判断是哪一只个体，或判断是否为已知个体。",
    )
    add_bullet(doc, "问题形态：closed-set 个体分类、open-set 个体验证、少样本个体注册、跨地点/跨设备声纹检索。")
    add_bullet(doc, "和 BioDCASE 的对应：Task 1 提醒多设备录音要先解决同步/时钟漂移；Task 4 对应有限标注预算；Task 5 对应跨域泛化；Task 6 提醒从检测转向数量/个体层面的生态指标。")
    add_bullet(doc, "技术路线：先做 tiger/leopard vocalization detector，再做 individual embedding；训练可用 metric learning、prototype learning、speaker verification 式 loss，以及 open-set 阈值校准。")
    add_bullet(doc, "数据设计：每只个体至少需要多天、多地点或多设备录音；划分时按日期/地点留出，避免同一 session 泄漏导致虚高。")
    add_bullet(doc, "和 Bird-MAE 的结合：把 Bird-MAE/音频 MAE 表征作为初始化，比较 self-supervised embedding 在低样本个体识别上的提升；必要时加入物种内声纹对比学习。")

    add_heading(doc, "可选研究题目表述", 2)
    make_table(
        doc,
        ["方向", "推荐任务定义", "评价方式"],
        [
            ("虎/豹个体识别", "给定目标个体库和一段新声音，判断是哪一只个体。", "Top-1/Top-k accuracy、macro-F1；按个体均衡评估。"),
            ("开放集个体验证", "判断一段声音是否来自已登记个体；若不是，则标为 unknown。", "AUROC、EER、FPR@TPR；更贴近野外新增个体。"),
            ("少样本个体注册", "每只新个体只有少量样例，系统快速建立个体声纹。", "N-shot accuracy、检索 mAP、随样本数变化的学习曲线。"),
            ("跨设备/跨地点泛化", "训练设备和测试设备不同，或保护区不同。", "domain-heldout accuracy、性能下降幅度、校准误差。"),
        ],
        [1.25, 3.30, 1.95],
        small=8.6,
    )

    add_heading(doc, "5. 照片候选与来源", 1)
    photos = [
        ("东北虎近景", "Wikimedia Commons: Amur Tiger Panthera tigris altaica Eye", "Public domain", "https://commons.wikimedia.org/wiki/File:Amur_Tiger_Panthera_tigris_altaica_Eye_2112px_edit.jpg"),
        ("东北豹相机陷阱", "Wikimedia Commons: Amur leopard. Frame from a camera trap", "Attribution", "https://commons.wikimedia.org/wiki/File:Amur_leopard._Frame_from_a_camera_trap.jpg"),
        ("东北豹正面照", "Wikimedia Commons: Amur Leopard Panthera pardus orientalis Facing Forward", "CC BY-SA 2.5", "https://commons.wikimedia.org/wiki/File:Amur_Leopard_Panthera_pardus_orientalis_Facing_Forward_1761px.jpg"),
    ]
    for title, desc, lic, url in photos:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.36)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        font(p.add_run(f"{title}：{desc}（{lic}） "), size=10.1)
        add_link(p, "链接", url)

    add_heading(doc, "6. 主要来源", 1)
    sources = [
        ("BioDCASE Challenge 2025", "https://biodcase.github.io/challenge2025/"),
        ("2025 Task 1 Multi-Channel Alignment", "https://biodcase.github.io/challenge2025/task1"),
        ("2025 Task 2 Whale Calls", "https://biodcase.github.io/challenge2025/task2"),
        ("2025 Task 3 Tiny Hardware", "https://biodcase.github.io/challenge2025/task3"),
        ("BioDCASE Challenge 2026", "https://biodcase.github.io/challenge2026/"),
        ("2026 Task 1 Multichannel Alignment", "https://biodcase.github.io/challenge2026/task1"),
        ("2026 Task 2 Whale Calls", "https://biodcase.github.io/challenge2026/task2"),
        ("2026 Task 3 Tiny Hardware", "https://biodcase.github.io/challenge2026/task3"),
        ("2026 Task 4 Active Learning", "https://biodcase.github.io/challenge2026/task4"),
        ("2026 Task 5 Cross-Domain Mosquito Species Classification", "https://biodcase.github.io/challenge2026/task5"),
        ("2026 Task 6 Bird Counting", "https://biodcase.github.io/challenge2026/task6"),
        ("Workshop 2025", "https://biodcase.github.io/workshop2025/"),
        ("Workshop 2026", "https://biodcase.github.io/workshop2026/"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.36)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        add_link(p, label, url)

    doc.core_properties.title = "BioDCASE 2025-2026 动物声学比赛总结更新版"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
